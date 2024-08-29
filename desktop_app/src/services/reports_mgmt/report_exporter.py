import os
from datetime import datetime
from typing import Dict, Any

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill


class ReportExporter:
    def __init__(self, report_data: Dict[str, Any], report_type: str, template_dir: str = './resources/templates'):
        self.report_data = report_data
        self.report_type = report_type
        self.template_dir = template_dir

    def export(self, output_filename: str):
        if self.report_type == 'docx':
            self._export_docx(output_filename)
        elif self.report_type == 'excel':
            self._export_excel(output_filename)
        else:
            raise ValueError(f"Unsupported report type: {self.report_type}")

    def _export_docx(self, output_filename: str):
        template_path = os.path.join(self.template_dir, f"{self.report_data['report_type']}_template.docx")

        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()

        # Add title
        doc.add_heading(self.report_data['title'], 0)

        # Add date range if available
        if 'date_range' in self.report_data:
            doc.add_paragraph(f"Date Range: {self.report_data['date_range']}")

        # Add summary if available
        if 'summary' in self.report_data:
            doc.add_paragraph(self.report_data['summary'])

        # Add table if available
        if 'table_data' in self.report_data:
            table = doc.add_table(rows=1, cols=len(self.report_data['table_headers']))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(self.report_data['table_headers']):
                hdr_cells[i].text = header

            for row in self.report_data['table_data']:
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

        doc.save(output_filename)

    def _export_excel(self, output_filename: str):
        template_path = os.path.join(self.template_dir, f"{self.report_data['report_type']}_template.xlsx")

        if os.path.exists(template_path):
            wb = load_workbook(template_path)
            sheet = wb.active
        else:
            wb = Workbook()
            sheet = wb.active

        sheet.title = self.report_data['title']

        # Add title
        sheet['A1'] = self.report_data['title']
        sheet['A1'].font = Font(size=16, bold=True)
        sheet['A1'].alignment = Alignment(horizontal='center')

        current_row = 2

        # Add date range if available
        if 'date_range' in self.report_data:
            sheet.cell(row=current_row, column=1, value=f"Date Range: {self.report_data['date_range']}")
            current_row += 1

        # Add summary if available
        if 'summary' in self.report_data:
            sheet.cell(row=current_row, column=1, value=self.report_data['summary'])
            current_row += 1

        # Add table if available
        if 'table_data' in self.report_data:
            for col, header in enumerate(self.report_data['table_headers'], start=1):
                cell = sheet.cell(row=current_row, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            current_row += 1

            for row in self.report_data['table_data']:
                for col, value in enumerate(row, start=1):
                    sheet.cell(row=current_row, column=col, value=value)
                current_row += 1

        wb.save(output_filename)


def format_report_data(report: Any) -> Dict[str, Any]:
    formatted_data = {
        'report_type': type(report).__name__,
        'title': type(report).__name__.replace('Report', ' Report').replace('Summary', ' Summary'),
    }

    if hasattr(report, 'start_date') and hasattr(report, 'end_date'):
        start_date = datetime.fromtimestamp(report.start_date).strftime('%Y-%m-%d')
        end_date = datetime.fromtimestamp(report.end_date).strftime('%Y-%m-%d')
        formatted_data['date_range'] = f"{start_date} to {end_date}"

    if hasattr(report, 'summary'):
        formatted_data['summary'] = str(report.summary)

    if hasattr(report, 'items'):
        formatted_data['table_headers'] = ['Product ID', 'Product Name', 'Quantity', 'Value']
        formatted_data['table_data'] = [
            [item.product_id, item.product_name, item.quantity, f"${item.value:.2f}"]
            for item in report.items
        ]
    elif hasattr(report, 'metrics'):
        formatted_data['table_headers'] = ['Metric', 'Value', 'Unit']
        formatted_data['table_data'] = [
            [metric.name, metric.value, metric.unit]
            for metric in report.metrics
        ]
    elif hasattr(report, 'products'):
        formatted_data['table_headers'] = ['Product ID', 'Product Name', 'Total Sold', 'Revenue']
        formatted_data['table_data'] = [
            [product.product_id, product.product_name, product.total_sold, f"${product.revenue:.2f}"]
            for product in report.products
        ]
    elif hasattr(report, 'suppliers'):
        formatted_data['table_headers'] = ['Supplier ID', 'Supplier Name', 'On-Time Delivery Rate', 'Quality Rating',
                                           'Total Orders']
        formatted_data['table_data'] = [
            [supplier.supplier_id, supplier.supplier_name, f"{supplier.on_time_delivery_rate:.2f}%",
             f"{supplier.quality_rating:.2f}", supplier.total_orders]
            for supplier in report.suppliers
        ]
    elif hasattr(report, 'movements'):
        formatted_data['table_headers'] = ['Product ID', 'Product Name', 'Quantity Change', 'Movement Type', 'Date']
        formatted_data['table_data'] = [
            [movement.product_id, movement.product_name, movement.quantity_change,
             movement.movement_type, datetime.fromtimestamp(movement.date).strftime('%Y-%m-%d %H:%M:%S')]
            for movement in report.movements
        ]
    elif hasattr(report, 'pickers'):
        formatted_data['table_headers'] = ['User ID', 'User Name', 'Total Picks', 'Average Pick Time', 'Accuracy Rate']
        formatted_data['table_data'] = [
            [picker.user_id, picker.user_name, picker.total_picks,
             f"{picker.average_pick_time:.2f} min", f"{picker.accuracy_rate:.2f}%"]
            for picker in report.pickers
        ]
    elif hasattr(report, 'zones'):
        formatted_data['table_headers'] = ['Zone ID', 'Zone Name', 'Total Capacity', 'Used Capacity',
                                           'Utilization Rate']
        formatted_data['table_data'] = [
            [zone.zone_id, zone.zone_name, zone.total_capacity,
             zone.used_capacity, f"{zone.utilization_rate:.2f}%"]
            for zone in report.zones
        ]
    elif hasattr(report, 'reasons'):
        formatted_data['table_headers'] = ['Reason', 'Count', 'Percentage']
        formatted_data['table_data'] = [
            [reason.reason, reason.count, f"{reason.percentage:.2f}%"]
            for reason in report.reasons
        ]
        formatted_data['summary'] = f"Total Returns: {report.total_returns}\nReturn Rate: {report.return_rate:.2f}%"

    return formatted_data
